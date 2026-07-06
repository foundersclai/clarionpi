"""Pure artifact builders — ``letter.docx`` + ``chronology.xlsx`` (package_builder / inv 10, 11).

Two deterministic, side-effect-free byte builders. Each takes already-approved state (rendered
section previews / rendered chronology rows) and returns the artifact bytes; neither touches the
DB, the registry, or storage — orchestration lives in :mod:`app.package.build`.

Two invariants are *this module*:

* **[11] Zero tokens in artifacts.** Every string that lands in the document is scanned with the
  registry's :data:`~app.engine.tokenizer.registry.TOKEN_RE` before the file is finalized; a hit
  raises :class:`ArtifactTokenLeak` (the build fails loud rather than shipping a token-shaped
  string in an attorney deliverable). The upstream render should have resolved every token — a
  leak here is a data bug, exactly as :mod:`app.api.wire_guard` treats a wire leak.
* **[10] Deterministic bytes.** No wall-clock enters the artifact bytes: the docx core properties
  and the xlsx workbook properties are pinned to a fixed timestamp (:data:`_PINNED_TS`), so
  identical inputs always produce an identical sha256. (The ``ArtifactSet`` row's ``created_at``
  is a separate, wall-clock DB default — it is not in the bytes.)

``python-docx`` / ``openpyxl`` are imported inside the functions to keep the module import-light,
mirroring :mod:`tests.corpus.pdf_builders`.
"""

from __future__ import annotations

import io
from collections.abc import Sequence
from datetime import datetime

from app.engine.tokenizer.registry import TOKEN_RE
from app.models.orm import DraftSection

# Fixed metadata timestamp stamped into the docx/xlsx bytes so the output is byte-deterministic
# (inv 10). Deliberately NOT ``now()`` — a wall-clock here would change the sha256 every build.
_PINNED_TS = datetime(2026, 1, 1)

# The author/creator recorded in the pinned document metadata.
_AUTHOR = "ClarionPI"

# The chronology worksheet header row (the ``render_rows_for_wire`` dict keys drive the cells).
_CHRONOLOGY_HEADERS = ("Date of service", "Provider", "Facility", "Type", "Narrative")


class ArtifactTokenLeak(Exception):
    """A token-shaped string reached an artifact — a resolution bug, never user error (inv 11).

    Carries the ``section_id`` (or a synthetic locus like ``"chronology"``) and the offending
    ``token`` so the failure names the leak precisely, mirroring
    :class:`app.api.wire_guard.TokenLeak`.
    """

    def __init__(self, section_id: str, token: str) -> None:
        self.section_id = section_id
        self.token = token
        super().__init__(f"token-shaped string {token!r} in artifact at section {section_id!r}")


def _scan_or_raise(text: str, *, section_id: str) -> None:
    """Raise :class:`ArtifactTokenLeak` if ``text`` carries anything token-shaped (inv 11)."""
    match = TOKEN_RE.search(text)
    if match is not None:
        raise ArtifactTokenLeak(section_id, match.group(0))


def _section_heading(section_id: str) -> str:
    """A human heading from a ``section_id`` — underscores to spaces, title-cased.

    e.g. ``"liability_summary"`` -> ``"Liability Summary"``. Deterministic and content-free (no
    registry lookup), so it never introduces a token.
    """
    return section_id.replace("_", " ").strip().title()


def build_letter_docx(
    *,
    firm_name: str,
    client_display_name: str,
    sections: Sequence[DraftSection],
    memo: str | None = None,
) -> bytes:
    """Build the demand ``letter.docx`` from a matter's PASSED, rendered sections.

    Layout (v1 — a *generated* letterhead; template ingestion is a recorded open question, so the
    letterhead is the firm name heading + a horizontal rule, not an uploaded firm template):

    * a level-1 heading with ``firm_name`` + a rule paragraph (the generated letterhead slot);
    * a ``Re:`` line naming ``client_display_name``;
    * per section (in the order given — the caller sorts by ``sort_order``): a heading derived from
      ``section_id`` (:func:`_section_heading`) followed by the section's ``rendered_preview`` split
      into paragraphs on blank lines (``\\n\\n``).

    The ``memo`` is deliberately NOT written into the letter: the strategy memo is a *separate*
    attorney-visible matter artifact (never sent to the carrier), so v1 excludes it from the
    carrier-facing letter even when one is supplied. The parameter is accepted for a stable
    signature; a future wave may emit it as its own artifact.

    Every paragraph's text is token-scanned before the file is finalized (:func:`_scan_or_raise`);
    a token-shaped string raises :class:`ArtifactTokenLeak` (inv 11). Core properties are pinned to
    :data:`_PINNED_TS` for byte determinism (inv 10). Returns the ``.docx`` bytes.
    """
    from docx import Document

    document = Document()

    # -- generated letterhead (template ingestion is a recorded open question; v1 = generated) --
    document.add_heading(firm_name, level=1)
    document.add_paragraph("—" * 24)  # a horizontal rule of em-dashes (deterministic)
    document.add_paragraph(f"Re: {client_display_name}")

    for section in sections:
        heading = _section_heading(section.section_id)
        document.add_heading(heading, level=2)
        preview = section.rendered_preview or ""
        # Split on blank lines into paragraphs; keep each paragraph's internal single newlines
        # collapsed to spaces so a paragraph is one docx paragraph.
        for block in preview.split("\n\n"):
            para_text = " ".join(line.strip() for line in block.splitlines()).strip()
            if not para_text:
                continue
            _scan_or_raise(para_text, section_id=section.section_id)
            document.add_paragraph(para_text)

    # -- pinned metadata for byte determinism (inv 10) --
    core = document.core_properties
    core.author = _AUTHOR
    core.created = _PINNED_TS
    core.modified = _PINNED_TS
    core.last_modified_by = _AUTHOR
    # Drop the auto-generated revision/last-printed so nothing wall-clock or counter-like leaks in.
    core.revision = 1

    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def build_chronology_xlsx(rows: Sequence[dict]) -> bytes:
    """Build the ``chronology.xlsx`` from ``render_rows_for_wire`` dicts (already token-free).

    One header row (:data:`_CHRONOLOGY_HEADERS`) then one row per input dict, reading
    ``date_of_service`` / ``provider_display`` / ``facility_display`` / ``encounter_type`` /
    ``narrative``. The narrative cell is token-scanned (:func:`_scan_or_raise`) — the rows come from
    ``resolve_text_for_wire`` which already asserts token-free, but the artifact layer re-checks
    (defence in depth, inv 11). Column widths are set to sane fixed values and the workbook
    properties are pinned to :data:`_PINNED_TS` for byte determinism (inv 10). Returns the
    ``.xlsx`` bytes.
    """
    from openpyxl import Workbook

    workbook = Workbook()
    worksheet = workbook.active
    worksheet.title = "Chronology"
    worksheet.append(list(_CHRONOLOGY_HEADERS))

    for index, row in enumerate(rows):
        narrative = str(row.get("narrative", ""))
        _scan_or_raise(narrative, section_id=f"chronology:{index}")
        worksheet.append(
            [
                str(row.get("date_of_service", "")),
                str(row.get("provider_display", "")),
                str(row.get("facility_display", "")),
                str(row.get("encounter_type", "")),
                narrative,
            ]
        )

    # Sane fixed column widths (deterministic; no autosize wall-clock/rendering dependence).
    for column_letter, width in (("A", 16), ("B", 28), ("C", 28), ("D", 12), ("E", 80)):
        worksheet.column_dimensions[column_letter].width = width

    # -- pinned metadata for byte determinism (inv 10) --
    props = workbook.properties
    props.creator = _AUTHOR
    props.lastModifiedBy = _AUTHOR
    props.created = _PINNED_TS
    props.modified = _PINNED_TS

    buffer = io.BytesIO()
    workbook.save(buffer)
    return buffer.getvalue()


__all__ = [
    "ArtifactTokenLeak",
    "build_chronology_xlsx",
    "build_letter_docx",
]
