"""THE M5 MILESTONE EXIT — one readable end-to-end demand build over HTTP.

Exit criterion (M5 plan): an end-to-end demand on a fixture matter runs the whole gate arc
``G1 → G3`` and produces ``letter.docx`` + ``binder.pdf`` (and the chronology + provenance
artifacts) with ZERO unresolved tokens and ledger-exact numbers. The arc runs through the REAL
wire (session-mode cookie login, the gates envelope, the M5 drafting/compliance/package routes,
typed refusals) exactly as the frontend drives it: Brain-2 drafts tokenized sections, the
compliance panel grades them, G3 approves, and the package builder emits the four immutable
artifacts — then a drift tripwire proves the built package is immutable.

The flow:

1. create the matter (attorney) → scripted Phase 0 over HTTP (``ingest/run`` SSE): ONE medical doc
   (two encounters 47 days apart → a >30-day treatment gap) + ONE bill doc with a KNOWN cents value,
   driven by a ``ScriptedProvider`` + ``FakeOcr`` + a tmp-dir ``LocalDiskStorage``. Reuses the
   M4-exit scripted-phase0 pattern wholesale (classify-then-window per doc, docs in
   ``(created_at, id)`` order).
2. G1 (facts_review): confirm every deadline candidate, approve → strategy_intake.
3. G1.5 (strategy_intake): anchor_amount + MMI after both encounters (a post-MMI gap, not adverse),
   approve → analysis_running.
4. analysis run over HTTP SSE (``NullProvider`` — deterministic flags only; ≥1 HIGH treatment gap).
5. G2a prep (attorney): disposition the HIGH gap flag (``address_in_letter``), pick the medical
   doc's exhibit + clear its PHI, then G2a approve → plan_review + the RegistryVersion frozen.
6. G2.5: POST plan/emit (``NullProvider`` → emphasis ``[]``); read the plan; assert the
   deterministic allocator gave the damages/demand sections the AMT ids; G2.5 approve → drafting.
7. drafting: build a ``ScriptedProvider`` script AFTER reading the plan (memo + per-section bodies
   composed from each section's required/allowed lists — tokens only, no literal dollars, under the
   word caps — + one clean judge reply per section). POST demand/generate → assert SSE order (memo
   step, five section frames with non-empty rendered previews, gate_ready compliance_review), draft
   VALIDATED → IN_COMPLIANCE, zero open blocking findings, a compliance_pass audit.
8. G3 approve → package_assembly → POST package/build → four artifact_ready frames → package_ready.
9. THE EXIT ASSERTIONS: download all four artifacts over HTTP; ``letter.docx`` opens via python-docx
   and its full text has ZERO ``TOKEN_RE`` matches AND contains the ledger's grand-billed display
   form EXACTLY (``cents_to_display`` of the known fixture cents — "ledger-exact numbers");
   ``binder.pdf`` opens via pypdf with the expected page count + outline; the provenance report's
   fact entries == the total rendered spans across the draft's sections; ``chronology.xlsx`` opens;
   every ``DraftSection`` validation == passed; the matter's final state is package_ready.
10. a drift tripwire: after package_ready, mint an attorney fact (a registry bump) → assert the
    machine refuses ``REGISTRY_BUMPED`` at package_ready (``IllegalTransition`` — immutable).

Prints a compact exit trail at the end (states → records → sections → findings → artifacts w/ shas).
Synthetic data only — no PHI.
"""

from __future__ import annotations

import datetime as dt
import io
import json
import re
import uuid
from collections.abc import Iterator
from pathlib import Path

import pytest
from fastapi.testclient import TestClient
from sqlalchemy import func, select
from sqlalchemy.orm import Session, sessionmaker

from app.api.deps import (
    DEV_USER_EMAIL,
    DEV_USER_PASSWORD,
    seed_dev_users,
)
from app.api.routes.ingest import get_ocr, get_provider
from app.api.routes.uploads import get_object_storage
from app.core.config import get_settings
from app.core.llm_provider import CompletionResult, NullProvider, ScriptedProvider
from app.core.storage import LocalDiskStorage
from app.core.tenancy import tenant_add
from app.corpus.ocr import FakeOcr
from app.engine.orchestrator.errors import IllegalTransition
from app.engine.orchestrator.machine import advance
from app.engine.tokenizer.registry import TOKEN_RE
from app.main import app
from app.models.enums import (
    ArtifactKind,
    DedupStatus,
    DocStatus,
    DocType,
    DraftStatus,
    FlagDisposition,
    FlagKind,
    FlagSeverity,
    GateEvent,
    GateState,
    LedgerCategory,
    PhiDisposition,
    SectionValidation,
)
from app.models.orm import (
    ArtifactSet,
    AuditEvent,
    BillingLine,
    CaseDocument,
    ComplianceFinding,
    DemandDraft,
    DraftSection,
    Matter,
    MedicalEncounter,
    RegistryVersion,
    RiskFlag,
    StrategyPlan,
    User,
)
from app.money.types import cents_to_display
from tests.corpus.pdf_builders import build_text_pdf

# Fixture dates: two encounters 47 days apart (a >30-day treatment gap), MMI AFTER both (so the gap
# counts and is a non-adverse post-MMI gap). The bill's billed cents are KNOWN so the letter's
# grand-billed display form is asserted exactly (ledger-exact numbers).
_ENC1_DOS = dt.date(2026, 2, 1)
_ENC2_DOS = dt.date(2026, 3, 20)  # 47 days after _ENC1_DOS
_MMI_DATE = "2026-04-01"  # after _ENC2_DOS
_BILL_CENTS = 250_000  # "$2,500.00" — the grand billed the letter must render exactly


# --------------------------------------------------------------------------------------
# Fixtures — mirror the M4-exit / analysis harness
# --------------------------------------------------------------------------------------


@pytest.fixture
def seeded(session_factory: sessionmaker[Session]) -> sessionmaker[Session]:
    db = session_factory()
    try:
        seed_dev_users(db)
    finally:
        db.close()
    return session_factory


@pytest.fixture
def session_mode(monkeypatch: pytest.MonkeyPatch) -> Iterator[None]:
    monkeypatch.setenv("AUTH_MODE", "session")
    get_settings.cache_clear()
    try:
        yield
    finally:
        get_settings.cache_clear()


@pytest.fixture
def phase0_overrides(tmp_path: Path) -> Iterator[LocalDiskStorage]:
    """Override ingest's storage/OCR deps with a tmp-dir store + FakeOcr (provider is per-call).

    The SAME storage backs Phase 0, the package build's binder collation, and the artifact
    downloads — so the PDF bytes the seed writes are the bytes the binder reads and the download
    route serves. The provider is per-phase (scripted phase0, null analysis/emit, scripted
    drafting), so it is set by each phase, never here.
    """
    storage = LocalDiskStorage(tmp_path / "storage")
    app.dependency_overrides[get_object_storage] = lambda: storage
    app.dependency_overrides[get_ocr] = FakeOcr
    try:
        yield storage
    finally:
        app.dependency_overrides.pop(get_object_storage, None)
        app.dependency_overrides.pop(get_ocr, None)


# --------------------------------------------------------------------------------------
# Small helpers — wire + scripted model replies
# --------------------------------------------------------------------------------------


def _login(client: TestClient, email: str) -> None:
    resp = client.post("/api/auth/login", json={"email": email, "password": DEV_USER_PASSWORD})
    assert resp.status_code == 200, resp.text


def _current(client: TestClient, matter_id: uuid.UUID) -> dict:
    resp = client.get(f"/api/matters/{matter_id}/gates/current")
    assert resp.status_code == 200, resp.text
    return resp.json()


def _submit(client: TestClient, matter_id: uuid.UUID, gate: str, payload: dict) -> object:
    return client.post(f"/api/matters/{matter_id}/gates/{gate}/submit", json=payload)


def _sse(resp_text: str) -> list[tuple[str, dict]]:
    """Parse the SSE text into ``(event, data)`` pairs (data JSON-decoded)."""
    parsed: list[tuple[str, dict]] = []
    for frame in resp_text.split("\n\n"):
        frame = frame.strip()
        if not frame.startswith("event: "):
            continue
        lines = frame.split("\n")
        event = lines[0].removeprefix("event: ")
        data = json.loads(lines[1].removeprefix("data: ")) if len(lines) > 1 else {}
        parsed.append((event, data))
    return parsed


def _result(text: str) -> CompletionResult:
    return CompletionResult(text=text, input_tokens=20, output_tokens=10, cost_cents=1)


def _classify(doc_type: str) -> CompletionResult:
    return _result(json.dumps({"doc_type": doc_type, "confidence": 0.95, "rationale": "r"}))


def _two_encounters(page: int) -> CompletionResult:
    """A medical-window reply carrying TWO encounters (different DOS → they never merge)."""
    return _result(
        json.dumps(
            {
                "encounters": [
                    {
                        "date_of_service": _ENC1_DOS.isoformat(),
                        "provider": "Dr. Erner",
                        "facility": "Mercy General",
                        "encounter_type": "ER",
                        "complaints": ["neck pain"],
                        "findings": ["tenderness"],
                        "diagnoses": ["cervical strain"],
                        "procedures": [],
                        "work_status": "light duty",
                        "anchor_pages": [page],
                        "field_confidence": {"provider": 0.6, "date_of_service": 0.9},
                    },
                    {
                        "date_of_service": _ENC2_DOS.isoformat(),
                        "provider": "Dr. Smith",
                        "facility": "Ortho Clinic",
                        "encounter_type": "office visit",
                        "complaints": ["residual stiffness"],
                        "findings": [],
                        "diagnoses": ["cervical strain"],
                        "procedures": [],
                        "work_status": "full duty",
                        "anchor_pages": [page],
                        "field_confidence": {"provider": 0.9, "date_of_service": 0.9},
                    },
                ]
            }
        )
    )


def _one_bill(page: int) -> CompletionResult:
    return _result(
        json.dumps(
            {
                "lines": [
                    {
                        "provider": "Mercy General",
                        "date_of_service": _ENC1_DOS.isoformat(),
                        "code": "99284",
                        "billed": cents_to_display(_BILL_CENTS),  # "$2,500.00" → exact cents
                        "adjusted": None,
                        "paid": None,
                        "outstanding": None,
                        "category": LedgerCategory.ER.value,
                        "anchor_page": page,
                    }
                ]
            }
        )
    )


def _make_uploaded_doc(
    db: Session,
    user: User,
    matter_id: uuid.UUID,
    storage: LocalDiskStorage,
    *,
    tag: str,
    filename: str,
) -> CaseDocument:
    """Store a 1-page synthetic PDF and create an UPLOADED, doc_type=other, unique CaseDocument."""
    key = f"matters/{matter_id}/{uuid.uuid4()}.pdf"
    storage.put(key, build_text_pdf([f"Record {tag}: patient reports neck pain after an MVA."]))
    doc = CaseDocument(
        matter_id=matter_id,
        doc_type=DocType.OTHER.value,
        source_label=filename,
        filename=filename,
        storage_key=key,
        dedup_status=DedupStatus.UNIQUE.value,
        status=DocStatus.UPLOADED.value,
    )
    tenant_add(db, doc, user.firm_id)
    db.commit()
    return doc


def _compose_section_body(section: dict) -> str:
    """Compose a tokens-only drafted body from a plan section's required/allowed lists.

    Respects the section contract mechanically (so the deterministic validator passes on the first
    attempt — no content retry, so the scripted script is exactly ``[memo, section×5, judge×5]``):

    * uses ONLY tokens from ``allowed_tokens`` (bracketed); INCLUDES every ``required_tokens`` id;
    * writes NO literal dollar figures / names / dates (inv 3/5 — the validator rejects a literal
      ``$`` and any unregistered/disallowed token);
    * an ``intro_and_representation`` (no allowed tokens) gets token-free prose;
    * stays well under ``max_words``.
    """
    required = list(section.get("required_tokens") or [])
    if required:
        toks = " ".join(f"[[{bare}]]" for bare in required)
        return f"This section rests on the following established facts: {toks}."
    # No required tokens: still keep it token-free (allowed may be non-empty, but nothing is
    # obligatory, so the safest clean body uses no tokens at all).
    return "We represent the claimant and set out the basis for this demand in the sections below."


# --------------------------------------------------------------------------------------
# THE M5 EXIT
# --------------------------------------------------------------------------------------


def test_m5_exit_full_demand_package(
    client: TestClient,
    seeded: sessionmaker[Session],
    session_mode: None,
    phase0_overrides: LocalDiskStorage,
) -> None:
    storage = phase0_overrides

    # ---- create the matter (attorney, session-mode cookie login) -------------------------
    _login(client, DEV_USER_EMAIL)
    created = client.post(
        "/api/matters",
        json={
            "client_display_name": "M5 Exit Client",
            "claim_type": "mva",
            "incident_date": "2026-01-15",
            "jurisdiction": "AZ",
            # WI-2: the four intake flags are REQUIRED; all-"no" is the in-box matter.
            "public_entity_involved": "no",
            "plaintiff_is_minor": "no",
            "wrongful_death": "no",
            "coverage_dispute": "no",
        },
    )
    assert created.status_code == 201, created.text
    matter_id = uuid.UUID(created.json()["id"])
    assert created.json()["gate_state"] == "corpus_processing"

    # ---- seed the two uploaded docs, then drive Phase 0 over HTTP (scripted) --------------
    db = seeded()
    try:
        attorney = db.execute(select(User).where(User.email == DEV_USER_EMAIL)).scalar_one()
        doc_med = _make_uploaded_doc(
            db, attorney, matter_id, storage, tag="MED", filename="records.pdf"
        )
        doc_bill = _make_uploaded_doc(
            db, attorney, matter_id, storage, tag="BILL", filename="bill.pdf"
        )
        ordered = sorted((doc_med, doc_bill), key=lambda d: (d.created_at, d.id))
        med_id, bill_id = doc_med.id, doc_bill.id
        script_order = [d.id for d in ordered]
    finally:
        db.close()

    def _script_for(doc_id: uuid.UUID) -> list[CompletionResult]:
        if doc_id == med_id:
            return [_classify("medical_record"), _two_encounters(1)]
        return [_classify("bill"), _one_bill(1)]

    phase0_script: list[CompletionResult] = []
    for doc_id in script_order:
        phase0_script += _script_for(doc_id)
    app.dependency_overrides[get_provider] = lambda: ScriptedProvider(phase0_script)
    try:
        ingest = client.post(f"/api/matters/{matter_id}/ingest/run")
    finally:
        app.dependency_overrides.pop(get_provider, None)
    assert ingest.status_code == 200, ingest.text
    assert "gate_ready" in [n for n, _ in _sse(ingest.text)]

    # Phase 0 landed: 2 encounters, 1 bill line at the known cents, gate at facts_review.
    db = seeded()
    try:
        assert db.get(CaseDocument, med_id).status == DocStatus.EXTRACTED.value
        assert db.get(CaseDocument, bill_id).status == DocStatus.EXTRACTED.value
        encounters = list(
            db.execute(
                select(MedicalEncounter).where(MedicalEncounter.matter_id == matter_id)
            ).scalars()
        )
        assert len(encounters) == 2
        bills = list(
            db.execute(select(BillingLine).where(BillingLine.matter_id == matter_id)).scalars()
        )
        assert len(bills) == 1 and bills[0].billed_cents == _BILL_CENTS  # exact cents, no float
    finally:
        db.close()

    # ---- G1 (facts_review): confirm every candidate, approve → strategy_intake -----------
    envelope = _current(client, matter_id)
    assert envelope["gate"] == "facts_review"
    candidates = envelope["view_model"]["deadline_candidates"]
    confirmations = [{"rule_id": c["rule_id"], "confirmed": True} for c in candidates]
    edited = _submit(
        client,
        matter_id,
        "facts_review",
        {
            "action": "edit",
            "idempotency_key": "m5-confirm-deadlines",
            "payload_version": envelope["payload_version"],
            "edits": {"deadline_confirmations": confirmations},
        },
    )
    assert edited.status_code == 200, edited.text
    envelope = _current(client, matter_id)
    approved = _submit(
        client,
        matter_id,
        "facts_review",
        {
            "action": "approve",
            "idempotency_key": "m5-g1-approve",
            "payload_version": envelope["payload_version"],
        },
    )
    assert approved.status_code == 200, approved.text
    assert approved.json()["result"]["to_state"] == "strategy_intake"

    # ---- G1.5 (strategy_intake): anchor amount + MMI after both encounters ----------------
    envelope = _current(client, matter_id)
    assert envelope["gate"] == "strategy_intake"
    g15 = _submit(
        client,
        matter_id,
        "strategy_intake",
        {
            "action": "approve",
            "idempotency_key": "m5-g15-submit",
            "payload_version": envelope["payload_version"],
            "edits": {
                "liability_theory": "Rear-end, admitted liability.",
                "injury_framing": "Cervical strain; conservative care.",
                "anchor_amount_cents": 4_000_000,
                "mmi_date": _MMI_DATE,
            },
        },
    )
    assert g15.status_code == 200, g15.text
    assert g15.json()["result"]["to_state"] == "analysis_running"

    # ---- analysis run over HTTP SSE (NullProvider → deterministic flags only) -------------
    app.dependency_overrides[get_provider] = lambda: NullProvider()
    try:
        run = client.post(f"/api/matters/{matter_id}/analysis/run")
    finally:
        app.dependency_overrides.pop(get_provider, None)
    assert run.status_code == 200, run.text
    assert "gate_ready" in [n for n, _ in _sse(run.text)]

    # A HIGH treatment-gap flag (undispositioned) is present — the G2a confirm blocks on it.
    db = seeded()
    try:
        gap = db.execute(
            select(RiskFlag).where(
                RiskFlag.matter_id == matter_id,
                RiskFlag.kind == FlagKind.TREATMENT_GAP.value,
            )
        ).scalar_one()
        assert gap.severity == FlagSeverity.HIGH.value
        gap_flag_id = gap.id
    finally:
        db.close()

    # ---- G2a prep (attorney): disposition the HIGH flag, pick + clear the exhibit ---------
    disp = client.put(
        f"/api/flags/{gap_flag_id}/disposition",
        json={"disposition": FlagDisposition.ADDRESS_IN_LETTER.value},
    )
    assert disp.status_code == 200, disp.text

    pick = client.put(
        f"/api/matters/{matter_id}/exhibits",
        json={"document_id": str(med_id), "include_pages": [1], "sort_order": 1},
    )
    assert pick.status_code == 200, pick.text
    exhibit_id = pick.json()["id"]
    phi = client.post(
        f"/api/exhibits/{exhibit_id}/phi", json={"disposition": PhiDisposition.CLEARED.value}
    )
    assert phi.status_code == 200, phi.text

    # ---- G2a approve → plan_review + the RegistryVersion frozen ---------------------------
    version = _current(client, matter_id)["payload_version"]
    g2a = _submit(
        client,
        matter_id,
        "evidence_review",
        {"action": "approve", "idempotency_key": "m5-g2a-approve", "payload_version": version},
    )
    assert g2a.status_code == 200, g2a.text
    assert g2a.json()["result"]["to_state"] == GateState.PLAN_REVIEW.value

    # ---- G2.5: plan emit (null provider → empty emphasis) + read the allocation -----------
    app.dependency_overrides[get_provider] = lambda: NullProvider()
    try:
        emit = client.post(f"/api/matters/{matter_id}/plan/emit")
    finally:
        app.dependency_overrides.pop(get_provider, None)
    assert emit.status_code == 200, emit.text
    plan_view = emit.json()["plan"]
    assert plan_view["approved"] is False
    assert plan_view["emphasis_directives"] == []  # null provider degrades emphasis visibly
    sections_by_id = {s["section_id"]: s for s in plan_view["sections"]}
    assert set(sections_by_id) == {
        "intro_and_representation",
        "liability",
        "injuries_and_treatment",
        "damages_and_specials",
        "demand_and_deadline",
    }
    # The deterministic allocator handed the damages/demand sections their AMT ids.
    amt_ids = _amt_token_ids(seeded, matter_id)
    damages_required = set(sections_by_id["damages_and_specials"]["required_tokens"])
    demand_required = set(sections_by_id["demand_and_deadline"]["required_tokens"])
    assert damages_required and damages_required <= amt_ids  # grand-billed + demand-basis AMTs
    assert demand_required and demand_required <= amt_ids  # demand-basis AMT

    # ---- G2.5 approve → drafting ----------------------------------------------------------
    version = _current(client, matter_id)["payload_version"]
    g25 = _submit(
        client,
        matter_id,
        "plan_review",
        {"action": "approve", "idempotency_key": "m5-g25-approve", "payload_version": version},
    )
    assert g25.status_code == 200, g25.text
    assert g25.json()["result"]["to_state"] == GateState.DRAFTING.value

    # ---- drafting: build the scripted script AFTER reading the plan -----------------------
    # Read the APPROVED plan's sections from the DB (skeleton order) and compose each section's
    # tokenized body from its required/allowed lists (tokens only, no literal dollars, under caps).
    db = seeded()
    try:
        plan = max(
            db.execute(select(StrategyPlan).where(StrategyPlan.matter_id == matter_id)).scalars(),
            key=lambda p: p.version,
        )
        ordered_sections = list(plan.sections)  # skeleton (list) order — the draft order
    finally:
        db.close()

    section_bodies = [_compose_section_body(s) for s in ordered_sections]
    n_sections = len(ordered_sections)
    drafting_script: list[CompletionResult] = [_result(json.dumps({"memo": "Straightforward."}))]
    drafting_script += [_result(json.dumps({"body_tokenized": body})) for body in section_bodies]
    # Every section PASSES deterministic validation on attempt 1 → the compliance pass then runs the
    # judge once per section (all clean → no findings, gate advances).
    drafting_script += [_result(json.dumps({"findings": []})) for _ in range(n_sections)]

    app.dependency_overrides[get_provider] = lambda: ScriptedProvider(drafting_script)
    try:
        generate = client.post(f"/api/matters/{matter_id}/demand/generate")
    finally:
        app.dependency_overrides.pop(get_provider, None)
    assert generate.status_code == 200, generate.text
    gen_events = _sse(generate.text)
    gen_names = [n for n, _ in gen_events]

    # SSE order: a memo step, then five section frames with non-empty rendered previews, then a
    # gate_ready compliance_review.
    memo_steps = [d for n, d in gen_events if n == "status" and d.get("step") == "memo"]
    assert memo_steps, "no memo step frame"
    section_frames = [d for n, d in gen_events if n == "section"]
    assert len(section_frames) == n_sections
    assert all((d.get("rendered_preview") or "").strip() for d in section_frames)
    assert "gate_ready" in gen_names
    gate_ready = [d for n, d in gen_events if n == "gate_ready"][0]
    assert gate_ready["gate"] == "compliance_review"
    # The memo step precedes every section frame in the stream.
    assert gen_names.index("status") < gen_names.index("section")

    # The draft went VALIDATED → IN_COMPLIANCE, zero open blocking findings, a compliance audit.
    db = seeded()
    try:
        matter = db.get(Matter, matter_id)
        assert matter.gate_state == GateState.COMPLIANCE_REVIEW.value
        draft = db.execute(
            select(DemandDraft).where(DemandDraft.matter_id == matter_id)
        ).scalar_one()
        assert draft.status == DraftStatus.IN_COMPLIANCE.value
        draft_id = draft.id
        open_blocking = db.execute(
            select(func.count())
            .select_from(ComplianceFinding)
            .where(
                ComplianceFinding.draft_id == draft.id,
                ComplianceFinding.severity == "blocking",
                ComplianceFinding.status.notin_(("re_verified", "dispositioned")),
            )
        ).scalar_one()
        assert open_blocking == 0
        pass_audits = list(
            db.execute(
                select(AuditEvent).where(
                    AuditEvent.firm_id == matter.firm_id,
                    AuditEvent.event_kind == "compliance_pass_completed",
                )
            ).scalars()
        )
        assert pass_audits, "no compliance_pass_completed audit"
        # Every section PASSED deterministic validation.
        sections = list(
            db.execute(select(DraftSection).where(DraftSection.draft_id == draft.id)).scalars()
        )
        assert len(sections) == n_sections
        assert all(s.validation == SectionValidation.PASSED.value for s in sections)
        total_spans = sum(len(s.spans or []) for s in sections)
    finally:
        db.close()

    # ---- G3 approve → package_assembly ----------------------------------------------------
    version = _current(client, matter_id)["payload_version"]
    g3 = _submit(
        client,
        matter_id,
        "compliance_review",
        {"action": "approve", "idempotency_key": "m5-g3-approve", "payload_version": version},
    )
    assert g3.status_code == 200, g3.text
    assert g3.json()["result"]["to_state"] == GateState.PACKAGE_ASSEMBLY.value

    # ---- package build → four artifact_ready frames + package_ready -----------------------
    build = client.post(f"/api/matters/{matter_id}/package/build")
    assert build.status_code == 200, build.text
    build_events = _sse(build.text)
    artifact_frames = [d for n, d in build_events if n == "artifact_ready"]
    assert len(artifact_frames) == 4
    assert {d["artifact_kind"] for d in artifact_frames} == {
        ArtifactKind.LETTER_DOCX.value,
        ArtifactKind.BINDER_PDF.value,
        ArtifactKind.CHRONOLOGY_XLSX.value,
        ArtifactKind.PROVENANCE_REPORT.value,
    }
    assert "gate_ready" in [n for n, _ in build_events]
    build_gate_ready = [d for n, d in build_events if n == "gate_ready"][0]
    assert build_gate_ready["gate"] == "package_ready"

    # ---- THE EXIT ASSERTIONS: download all four + open + check ----------------------------
    from docx import Document  # imported here (mirrors the artifact builder's local import)
    from openpyxl import load_workbook
    from pypdf import PdfReader

    listing = client.get(f"/api/matters/{matter_id}/artifacts")
    assert listing.status_code == 200, listing.text
    sets = listing.json()["sets"]
    assert len(sets) == 1
    artifact_set = sets[0]
    artifacts = artifact_set["artifacts"]
    assert all("object_key" not in a for a in artifacts)  # object_key never on the wire

    downloaded: dict[str, bytes] = {}
    for a in artifacts:
        dl = client.get(a["url"])
        assert dl.status_code == 200, dl.text
        downloaded[a["kind"]] = dl.content

    # letter.docx: opens via python-docx; full text has ZERO tokens AND the grand-billed display
    # form EXACTLY (ledger-exact numbers).
    document = Document(io.BytesIO(downloaded[ArtifactKind.LETTER_DOCX.value]))
    letter_text = "\n".join(p.text for p in document.paragraphs)
    assert TOKEN_RE.search(letter_text) is None  # zero unresolved tokens
    grand_billed_display = cents_to_display(_BILL_CENTS)  # "$2,500.00"
    assert grand_billed_display in letter_text, (
        f"letter is missing the ledger grand-billed display form {grand_billed_display!r}"
    )

    # binder.pdf: opens via pypdf; index page + the one included page = 2 pages; one outline entry.
    binder_reader = PdfReader(io.BytesIO(downloaded[ArtifactKind.BINDER_PDF.value]))
    assert len(binder_reader.pages) == 2  # index page + 1 collated medical page
    assert len(binder_reader.outline) == 1  # one bookmark for the single exhibit

    # provenance report: fact entries == the total rendered spans across the sections. The report
    # prints "Total rendered facts: N"; re-derive N honestly from the DB spans and check equality.
    provenance_reader = PdfReader(io.BytesIO(downloaded[ArtifactKind.PROVENANCE_REPORT.value]))
    provenance_text = "\n".join(page.extract_text() or "" for page in provenance_reader.pages)
    assert f"Total rendered facts: {total_spans}" in provenance_text
    # Cross-check the count against the per-span id lines the report prints ("[FACT_n]"/"[AMT_n]").
    printed_span_lines = len(re.findall(r"\[(?:FACT|AMT|CITE|EX)_\d+\]", provenance_text))
    assert printed_span_lines == total_spans

    # chronology.xlsx: opens via openpyxl (header + one row per encounter).
    workbook = load_workbook(io.BytesIO(downloaded[ArtifactKind.CHRONOLOGY_XLSX.value]))
    worksheet = workbook.active
    assert worksheet.max_row >= 3  # header + two encounters

    # Final state: package_ready; exactly one ArtifactSet row.
    db = seeded()
    try:
        matter = db.get(Matter, matter_id)
        assert matter.gate_state == GateState.PACKAGE_READY.value
        rows = list(
            db.execute(select(ArtifactSet).where(ArtifactSet.matter_id == matter_id)).scalars()
        )
        assert len(rows) == 1
        the_set = rows[0]
    finally:
        db.close()

    # ---- THE DRIFT TRIPWIRE: a bump at package_ready is refused by the machine (immutable) -
    # Mint an attorney fact — a real registry bump. The package_ready state has NO REGISTRY_BUMPED
    # edge (flow_04: immutable, new records start a new draft cycle), so the machine refuses.
    db = seeded()
    try:
        from app.engine.tokenizer import registry

        matter = db.get(Matter, matter_id)
        before_version = matter.registry_version
        attorney = (
            db.get(User, the_set.built_by)
            or db.execute(select(User).where(User.email == DEV_USER_EMAIL)).scalar_one()
        )
        registry.mint_attorney_fact(
            db, matter=matter, user=attorney, display_form="a late fact", value={"note": "late"}
        )
        db.refresh(matter)
        assert matter.registry_version == before_version + 1  # the bump really happened
    finally:
        db.close()

    # The machine refuses REGISTRY_BUMPED from package_ready — immutability proven at the machine.
    with pytest.raises(IllegalTransition):
        advance(GateState.PACKAGE_READY, GateEvent.REGISTRY_BUMPED)

    # ---- the compact M5-exit trail summary (the evidence the integrator pastes) -----------
    db = seeded()
    try:
        matter = db.get(Matter, matter_id)
        draft = db.get(DemandDraft, draft_id)
        sections = sorted(
            db.execute(select(DraftSection).where(DraftSection.draft_id == draft_id)).scalars(),
            key=lambda s: s.sort_order,
        )
        findings = list(
            db.execute(
                select(ComplianceFinding).where(ComplianceFinding.draft_id == draft_id)
            ).scalars()
        )
        frozen = (
            db.execute(
                select(RegistryVersion).where(
                    RegistryVersion.matter_id == matter_id, RegistryVersion.frozen.is_(True)
                )
            )
            .scalars()
            .first()
        )
        print(f"\nM5 EXIT TRAIL — matter {matter_id}")
        print(
            "  states: corpus_processing -> facts_review -> strategy_intake -> analysis_running"
            " -> evidence_review -> plan_review -> drafting -> compliance_review"
            " -> package_assembly -> package_ready"
        )
        print(
            f"  gate records: G1(edit+approve) G1.5(approve) G2a(approve) G2.5(approve)"
            f" G3(approve)  RegistryVersion frozen @ v{frozen.version if frozen else '—'}"
        )
        print(
            f"  draft v{draft.version} status={draft.status} (plan v{draft.strategy_plan_version},"
            f" registry v{draft.registry_version}); memo chars={len(draft.memo or '')}"
        )
        for s in sections:
            preview = (s.rendered_preview or "").replace("\n", " ")
            print(
                f"    section {s.section_id:<24} validation={s.validation}"
                f" spans={len(s.spans or [])}  preview='{preview[:56]}'"
            )
        print(f"  findings: {len(findings)} total; open-blocking after compliance={open_blocking}")
        print(f"  ledger grand-billed rendered in letter: {grand_billed_display} (exact cents)")
        for a in the_set.artifacts or []:
            print(
                f"    artifact {a['kind']:<20} bytes={a['byte_count']:>6}  sha={a['sha256'][:16]}…"
            )
        print(
            f"  binder pages={len(binder_reader.pages)} outline={len(binder_reader.outline)};"
            f" provenance rendered-facts={total_spans}; chronology rows={worksheet.max_row}"
        )
        print(
            "  drift tripwire: attorney-fact bump -> registry_version"
            f" {before_version}->{before_version + 1}; REGISTRY_BUMPED @ package_ready ->"
            " IllegalTransition (immutable)"
        )
    finally:
        db.close()


def _amt_token_ids(session_factory: sessionmaker[Session], matter_id: uuid.UUID) -> set[str]:
    """The bare ids of the matter's AMOUNT FactToken slots (the two always-on ledger AMTs)."""
    from app.models.enums import TokenKind
    from app.models.orm import FactToken

    db = session_factory()
    try:
        rows = db.execute(
            select(FactToken.token_id).where(
                FactToken.matter_id == matter_id,
                FactToken.kind == TokenKind.AMOUNT.value,
            )
        ).scalars()
        return set(rows)
    finally:
        db.close()
